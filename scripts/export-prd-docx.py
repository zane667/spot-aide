#!/usr/bin/env python3
"""把 Markdown 导出为中文 Word（默认 PRD；可用参数导出其它文档）。"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x4B, 0x4B, 0x4B)
HEADER_BG = "1A1A1A"
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT = "FFD100"
ROW_ALT = "FFF8EB"
CODE_BG = "F4F4F4"


def set_run_font(run, east_asia: str, ascii_font: str = "Calibri", size_pt: float | None = None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = INK


def set_paragraph_format(p, before=0, after=6, line=1.5, first_line=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line is not None:
        fmt.first_line_indent = Cm(first_line)


def shade_cell(cell, fill: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D0D0D0")
        borders.append(el)
    old = tc_pr.find(qn("w:tcBorders"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(borders)


def set_cell_margins(cell, cm=0.12):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(int(cm * 567)))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(mar)


def add_page_border_header_footer(doc: Document, header_left: str, header_right: str):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = hp.add_run(header_left)
    set_run_font(r1, "微软雅黑", "Calibri", 9, bold=True)
    r2 = hp.add_run(f"  ·  {header_right}")
    set_run_font(r2, "微软雅黑", "Calibri", 9)
    r2.font.color.rgb = MUTED
    hp.paragraph_format.space_after = Pt(4)

    # yellow accent line under header
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), ACCENT)
    pBdr.append(bottom)
    hp._p.get_or_add_pPr().append(pBdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(fp, before=4, after=0, line=1.0)
    run = fp.add_run("内部评审资料  ·  第 ")
    set_run_font(run, "微软雅黑", "Calibri", 9)
    run.font.color.rgb = MUTED

    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r_page = fp.add_run()
    set_run_font(r_page, "微软雅黑", "Calibri", 9)
    r_page.font.color.rgb = MUTED
    r_page._r.append(fld1)
    r_page._r.append(instr)
    r_page._r.append(fld2)

    run2 = fp.add_run(" 页")
    set_run_font(run2, "微软雅黑", "Calibri", 9)
    run2.font.color.rgb = MUTED


def unescape(text: str) -> str:
    return (
        text.replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&amp;", "&")
        .replace("\\<", "<")
    )


def add_runs_with_markup(paragraph, text: str, east="宋体", size=12, base_bold=False):
    text = unescape(text)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, east, size_pt=size, bold=base_bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, east, size_pt=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "微软雅黑", "Consolas", size_pt=size - 0.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(label)
            set_run_font(run, east, size_pt=size)
            run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
            run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, east, size_pt=size, bold=base_bold)


def add_heading_styled(doc, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_format(p, before=18, after=8, line=1.3)
        run = p.add_run(text)
        set_run_font(run, "微软雅黑", size_pt=16, bold=True)
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), ACCENT)
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
    elif level == 2:
        set_paragraph_format(p, before=12, after=6, line=1.3)
        run = p.add_run(text)
        set_run_font(run, "微软雅黑", size_pt=13.5, bold=True)
    else:
        set_paragraph_format(p, before=10, after=4, line=1.3)
        run = p.add_run(text)
        set_run_font(run, "微软雅黑", size_pt=12, bold=True)
    return p


def add_body(doc, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=8, line=1.5)
    add_runs_with_markup(p, text)
    return p


def add_quote(doc, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=6, line=1.5)
    p.paragraph_format.left_indent = Cm(0.6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), ACCENT)
    pBdr.append(left)
    pPr.append(pBdr)
    add_runs_with_markup(p, text)
    return p


def add_list_item(doc, text: str, ordered: bool, index: int | None = None):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=1, after=3, line=1.4)
    p.paragraph_format.left_indent = Cm(0.75)
    prefix = f"{index}. " if ordered else "• "
    run = p.add_run(prefix)
    set_run_font(run, "微软雅黑", size_pt=12, bold=ordered)
    add_runs_with_markup(p, text)
    return p


def add_checkbox(doc, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=1, after=3, line=1.4)
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run("☐  ")
    set_run_font(run, "微软雅黑", size_pt=12)
    add_runs_with_markup(p, text)
    return p


def add_code_block(doc, lines: list[str]):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=6, after=10, line=1.25)
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_BG)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    run = p.add_run("\n".join(lines))
    set_run_font(run, "微软雅黑", "Consolas", size_pt=10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    usable = Cm(16.0)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, before=2, after=2, line=1.15)
            value = unescape(row[j] if j < len(row) else "")
            add_runs_with_markup(p, value, east="微软雅黑" if i == 0 else "宋体", size=10.5, base_bold=(i == 0))
            if i == 0:
                for run in p.runs:
                    run.font.color.rgb = HEADER_FG
                    run.bold = True
                shade_cell(cell, HEADER_BG)
            elif i % 2 == 0:
                shade_cell(cell, ROW_ALT)
            set_cell_borders(cell)
            set_cell_margins(cell)
    # keep table from overflowing: set preferred width
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, before=0, after=8, line=1.0)


def parse_table(block_lines: list[str]) -> list[list[str]]:
    rows = []
    for line in block_lines:
        if re.match(r"^\s*\|?\s*:?-{3,}", line.replace("|", " | ")):
            # separator row like | --- | --- |
            if set(line.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")) == set():
                continue
        if re.match(r"^\s*\|?\s*-{3,}", line) and "---" in line and not re.search(r"[^\s|:\-]", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_cover(doc: Document, title: str, subtitle: str, meta: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=36, after=4, line=1.15)
    run = p.add_run(title)
    set_run_font(run, "微软雅黑", size_pt=28, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=10, line=1.15)
    run = p.add_run(subtitle)
    set_run_font(run, "微软雅黑", size_pt=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=4, line=1.2)
    run = p.add_run(meta)
    set_run_font(run, "微软雅黑", size_pt=11)
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=18, line=1.2)
    run = p.add_run("https://spot-aide.vercel.app")
    set_run_font(run, "Calibri", "Calibri", size_pt=10.5)
    run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)

    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=16, line=1.0)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT)
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def convert(md: str, doc: Document):
    lines = md.splitlines()
    i = 0
    skipped_title = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not skipped_title and stripped.startswith("# "):
            skipped_title = True
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        if stripped.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i].rstrip())
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, parse_table(block))
            continue

        if stripped.startswith("#### "):
            add_heading_styled(doc, stripped[5:], 3)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading_styled(doc, stripped[4:], 3)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading_styled(doc, stripped[3:], 2)
            i += 1
            continue
        if stripped.startswith("# "):
            add_heading_styled(doc, stripped[2:], 1)
            i += 1
            continue

        m_ol = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_ol:
            add_list_item(doc, m_ol.group(2), ordered=True, index=int(m_ol.group(1)))
            i += 1
            continue

        if stripped.startswith("- [ ] "):
            add_checkbox(doc, stripped[6:])
            i += 1
            continue

        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:], ordered=False)
            i += 1
            continue

        if stripped == ">" or stripped.startswith("> "):
            quoted = stripped[2:] if stripped.startswith("> ") else ""
            if quoted:
                add_quote(doc, quoted)
            i += 1
            continue

        add_body(doc, stripped)
        i += 1


def parse_args() -> argparse.Namespace:
    root = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(description="导出探店参谋 Markdown 为 Word")
    parser.add_argument("--src", type=Path, default=root / "docs" / "prd.md")
    parser.add_argument("--out", type=Path, default=Path.home() / "Downloads" / "探店参谋-PRD.docx")
    parser.add_argument("--header-left", default="探店参谋")
    parser.add_argument("--header-right", default="产品需求文档（PRD）  ·  V1.0")
    parser.add_argument("--cover-title", default="探店参谋")
    parser.add_argument("--cover-subtitle", default="产品需求文档（PRD）")
    parser.add_argument("--cover-meta", default="SpotAide  ·  V1.0  ·  已实现并可演示")
    return parser.parse_args()


def main():
    args = parse_args()
    src = args.src if args.src.is_absolute() else Path.cwd() / args.src
    text = src.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(12)
    styles._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_page_border_header_footer(doc, args.header_left, args.header_right)
    add_cover(doc, args.cover_title, args.cover_subtitle, args.cover_meta)
    convert(text, doc)

    out = args.out.expanduser()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(str(out))


if __name__ == "__main__":
    main()
